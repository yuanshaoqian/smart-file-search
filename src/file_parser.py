#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件解析器模块
支持多种文件格式的文本内容提取
"""

import os
import re
import chardet
import logging
from pathlib import Path
from typing import Optional, Dict, Any
from abc import ABC, abstractmethod
from loguru import logger

# 尝试导入可选依赖
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx 未安装，Word 文档支持已禁用")

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    logger.warning("openpyxl 未安装，Excel 文件支持已禁用")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False


class FileParserBase(ABC):
    """文件解析器基类"""
    
    def __init__(self):
        self.logger = logger.bind(module="parser")
    
    @abstractmethod
    def parse(self, file_path: str) -> Optional[str]:
        """
        解析文件并返回文本内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析出的文本内容，如果解析失败则返回 None
        """
        pass
    
    @abstractmethod
    def supports(self, file_path: str) -> bool:
        """
        检查是否支持该文件类型
        
        Args:
            file_path: 文件路径
            
        Returns:
            是否支持
        """
        pass


class TextFileParser(FileParserBase):
    """纯文本文件解析器"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {'.txt', '.md', '.py', '.java', '.cpp', '.h', 
                                     '.json', '.xml', '.yml', '.yaml', '.html', 
                                     '.htm', '.css', '.js', '.ts', '.sql', '.sh',
                                     '.bat', '.ps1', '.cfg', '.ini', '.conf'}
    
    def parse(self, file_path: str) -> Optional[str]:
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
            
            # 检测编码
            result = chardet.detect(raw_data)
            encoding = result['encoding'] or 'utf-8'
            confidence = result['confidence']
            
            if confidence < 0.7:
                self.logger.warning(f"编码检测置信度低 ({confidence:.2f})，文件: {file_path}")
            
            # 尝试解码
            try:
                content = raw_data.decode(encoding)
            except UnicodeDecodeError:
                # 尝试常见编码
                for enc in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
                    try:
                        content = raw_data.decode(enc)
                        encoding = enc
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    self.logger.error(f"无法解码文件: {file_path}")
                    return None
            
            # 清理文本：移除过多的空白字符
            content = re.sub(r'\r\n', '\n', content)  # 统一换行符
            content = re.sub(r'\t', ' ', content)     # 制表符转空格
            content = re.sub(r'[ \t]{2,}', ' ', content)  # 多个空格合并
            
            return content
            
        except Exception as e:
            self.logger.error(f"解析文本文件失败 {file_path}: {e}")
            return None
    
    def supports(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions


class WordDocumentParser(FileParserBase):
    """Word 文档解析器"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {'.docx'}
    
    def parse(self, file_path: str) -> Optional[str]:
        if not DOCX_AVAILABLE:
            self.logger.error(f"无法解析 Word 文档，python-docx 未安装: {file_path}")
            return None
        
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # 提取表格内容
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        tables_text.append(' | '.join(row_text))
            
            content = '\n'.join(paragraphs)
            if tables_text:
                content += '\n\n表格内容:\n' + '\n'.join(tables_text)
            
            return content if content.strip() else None
            
        except Exception as e:
            self.logger.error(f"解析 Word 文档失败 {file_path}: {e}")
            return None
    
    def supports(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions and DOCX_AVAILABLE


class ExcelFileParser(FileParserBase):
    """Excel 文件解析器"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = {'.xlsx', '.xls'}
    
    def parse(self, file_path: str) -> Optional[str]:
        if not OPENPYXL_AVAILABLE:
            self.logger.error(f"无法解析 Excel 文件，openpyxl 未安装: {file_path}")
            return None
        
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
            content_parts = []
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheet_content = []
                
                # 获取有数据的范围
                if ws.max_row > 0 and ws.max_column > 0:
                    for row in ws.iter_rows(min_row=1, max_row=min(ws.max_row, 1000), 
                                          min_col=1, max_col=min(ws.max_column, 50),
                                          values_only=True):
                        # 过滤空行
                        row_values = []
                        for cell in row:
                            if cell is not None:
                                # 处理不同数据类型
                                if isinstance(cell, str):
                                    row_values.append(cell.strip())
                                else:
                                    row_values.append(str(cell))
                        
                        if row_values:
                            sheet_content.append(' | '.join(row_values))
                
                if sheet_content:
                    content_parts.append(f"工作表: {sheet_name}")
                    content_parts.extend(sheet_content)
                    content_parts.append('')  # 空行分隔
            
            wb.close()
            
            content = '\n'.join(content_parts)
            return content if content.strip() else None
            
        except Exception as e:
            self.logger.error(f"解析 Excel 文件失败 {file_path}: {e}")
            return None
    
    def supports(self, file_path: str) -> bool:
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions and OPENPYXL_AVAILABLE


class FileParser:
    """统一文件解析器，根据文件类型分发给具体解析器"""
    
    def __init__(self):
        self.logger = logger.bind(module="file_parser")
        self.parsers = []
        self._register_parsers()
    
    def _register_parsers(self):
        """注册所有可用的解析器"""
        self.parsers.append(TextFileParser())
        
        if DOCX_AVAILABLE:
            self.parsers.append(WordDocumentParser())
        
        if OPENPYXL_AVAILABLE:
            self.parsers.append(ExcelFileParser())
        
        self.logger.info(f"已注册 {len(self.parsers)} 个文件解析器")
    
    def get_parser(self, file_path: str) -> Optional[FileParserBase]:
        """
        获取适合文件类型的解析器
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析器实例，如果不支持则返回 None
        """
        for parser in self.parsers:
            if parser.supports(file_path):
                return parser
        return None
    
    def parse(self, file_path: str) -> Optional[str]:
        """
        解析文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            文件内容文本，如果解析失败则返回 None
        """
        file_path = str(Path(file_path).resolve())
        
        # 检查文件是否存在
        if not Path(file_path).exists():
            self.logger.error(f"文件不存在: {file_path}")
            return None
        
        # 检查文件大小
        file_size = Path(file_path).stat().st_size
        if file_size > 100 * 1024 * 1024:  # 100MB
            self.logger.warning(f"文件过大 ({file_size/1024/1024:.1f}MB)，跳过: {file_path}")
            return None
        
        # 获取合适的解析器
        parser = self.get_parser(file_path)
        if parser is None:
            self.logger.warning(f"不支持的文件类型: {file_path}")
            return None
        
        try:
            self.logger.debug(f"解析文件: {file_path}")
            content = parser.parse(file_path)
            
            if content and len(content.strip()) > 0:
                # 添加文件元信息
                meta = f"文件: {Path(file_path).name}\n路径: {file_path}\n大小: {file_size} 字节\n"
                content = meta + "=" * 50 + "\n" + content
                return content
            else:
                self.logger.warning(f"文件内容为空: {file_path}")
                return None
                
        except Exception as e:
            self.logger.error(f"解析文件失败 {file_path}: {e}")
            return None
    
    def get_supported_extensions(self) -> set:
        """获取所有支持的扩展名"""
        extensions = set()
        for parser in self.parsers:
            if hasattr(parser, 'supported_extensions'):
                extensions.update(parser.supported_extensions)
        return extensions
    
    def is_supported(self, file_path: str) -> bool:
        """检查文件是否支持"""
        return self.get_parser(file_path) is not None


# 全局解析器实例
_parser_instance: Optional[FileParser] = None

def get_parser() -> FileParser:
    """获取全局文件解析器"""
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = FileParser()
    return _parser_instance


if __name__ == "__main__":
    # 测试解析器
    parser = get_parser()
    
    # 测试文件
    test_files = [
        "/tmp/test.txt",
        "/tmp/test.docx",
        "/tmp/test.xlsx",
    ]
    
    for test_file in test_files:
        if Path(test_file).exists():
            print(f"\n解析文件: {test_file}")
            content = parser.parse(test_file)
            if content:
                print(f"内容预览: {content[:200]}...")
            else:
                print("解析失败或不支持")
        else:
            print(f"测试文件不存在: {test_file}")